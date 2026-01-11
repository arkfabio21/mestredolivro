#!/usr/bin/env python3
"""
Insere uma tabela em um documento Word existente.

USAGE:
    python3 scripts/insert-table.py documento.docx --after "texto para localizar"
"""

import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Define cor de fundo de uma célula."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def insert_table_after_paragraph(doc_path, output_path, search_text, table_data):
    """Insere tabela após parágrafo que contém o texto especificado."""
    doc = Document(doc_path)

    # Encontrar o parágrafo
    target_paragraph = None
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            target_paragraph = para
            break

    if not target_paragraph:
        print(f"Texto '{search_text}' não encontrado no documento.")
        return False

    # Inserir parágrafo vazio e depois a tabela
    # Criar tabela
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Preencher tabela
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)

            # Formatação do texto
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'EB Garamond'
                    run.font.size = Pt(11)

                    # Header em negrito
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # Cor de fundo do header
            if i == 0:
                set_cell_shading(cell, "2C3E50")
            elif i % 2 == 0:
                set_cell_shading(cell, "F8F9FA")

    # Mover tabela para depois do parágrafo alvo
    target_paragraph._element.addnext(table._tbl)

    # Adicionar espaço antes e depois da tabela
    spacing_before = doc.add_paragraph()
    spacing_before._element.addnext(table._tbl)

    doc.save(output_path)
    print(f"✅ Tabela inserida em {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(description='Inserir tabela em documento Word')
    parser.add_argument('documento', help='Caminho do documento Word')
    parser.add_argument('--output', '-o', help='Caminho de saída (padrão: sobrescreve)')
    parser.add_argument('--after', required=True, help='Texto após o qual inserir a tabela')

    args = parser.parse_args()

    output = args.output or args.documento

    # Dados da tabela de métricas
    table_data = [
        ["Indicador", "Antes", "Depois", "Variação"],
        ["Acidentes com afastamento", "14/ano", "3/ano", "-78%"],
        ["Quase-acidentes reportados", "23/trimestre", "312/trimestre", "+1.256%"],
        ["Tempo de supervisão em campo", "8%", "34%", "+325%"],
        ["Custo com afastamentos", "R$ 340 mil", "R$ 67 mil", "-80%"],
        ["Turnover no primeiro ano", "18%", "11%", "-39%"]
    ]

    insert_table_after_paragraph(args.documento, output, args.after, table_data)


if __name__ == "__main__":
    main()
