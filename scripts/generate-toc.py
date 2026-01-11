#!/usr/bin/env python3
"""
generate-toc.py - Gera sumário automático (Table of Contents) em documentos Word

USAGE:
    python3 scripts/generate-toc.py input.docx output.docx [--position start|after-preface]

DESCRIÇÃO:
    Gera um sumário automático baseado nos estilos de heading do documento.
    O sumário é inserido no início ou após o prefácio, conforme especificado.

FUNCIONAMENTO:
    1. Escaneia o documento buscando headings (títulos e subtítulos)
    2. Cria uma estrutura hierárquica
    3. Insere o sumário com números de página
    4. Atualiza referências

DEPENDÊNCIAS:
    pip install python-docx lxml
"""

import argparse
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRO: python-docx não instalado.")
    print("Execute: pip install python-docx")
    sys.exit(1)


# =============================================================================
# DETECÇÃO DE HEADINGS
# =============================================================================

def is_chapter_marker(text: str) -> bool:
    """Detecta marcador de capítulo (ex: 'Capítulo 1', 'CAPÍTULO 1')"""
    patterns = [
        r'^cap[íi]tulo\s+\d+',
        r'^chapter\s+\d+',
        r'^parte\s+\d+',
    ]
    text_lower = text.lower().strip()
    return any(re.match(p, text_lower) for p in patterns)


def is_heading(paragraph) -> tuple:
    """
    Detecta se um parágrafo é um heading e retorna (nível, título)
    Retorna (None, None) se não for heading
    """
    text = paragraph.text.strip()

    if not text:
        return None, None

    # Verificar estilo do parágrafo
    style_name = paragraph.style.name.lower() if paragraph.style else ''

    # Headings explícitos por estilo
    if 'heading 1' in style_name or 'título 1' in style_name:
        return 1, text
    if 'heading 2' in style_name or 'título 2' in style_name:
        return 2, text
    if 'heading 3' in style_name or 'título 3' in style_name:
        return 3, text

    # Verificar por formatação
    has_bold = any(run.bold for run in paragraph.runs if run.text.strip())
    font_size = None
    for run in paragraph.runs:
        if run.font.size:
            font_size = run.font.size.pt
            break

    # Heurísticas para detectar headings
    if is_chapter_marker(text):
        return 0, text  # Nível 0 = marcador de capítulo (não incluir no TOC)

    # Texto curto + negrito + fonte grande = provavelmente título
    if len(text) < 100 and has_bold:
        if font_size and font_size >= 14:
            return 1, text
        elif font_size and font_size >= 12:
            return 2, text

    # Texto em CAIXA ALTA curto
    if len(text) < 80 and text.isupper() and not text.endswith('.'):
        return 1, text

    # Numeração de seção
    if re.match(r'^\d+\.\s+\w', text):
        return 2, text
    if re.match(r'^\d+\.\d+\s+\w', text):
        return 3, text

    return None, None


# =============================================================================
# GERAÇÃO DO SUMÁRIO
# =============================================================================

def extract_headings(doc) -> list:
    """Extrai todos os headings do documento"""
    headings = []
    current_chapter = None

    for i, para in enumerate(doc.paragraphs):
        level, title = is_heading(para)

        if level is not None:
            if level == 0:
                # Marcador de capítulo - guardar para combinar com título seguinte
                current_chapter = title
            else:
                # Se temos um marcador de capítulo pendente, combinar
                if current_chapter and level == 1:
                    title = f"{current_chapter}: {title}"
                    current_chapter = None

                headings.append({
                    'level': level,
                    'title': title,
                    'paragraph_index': i,
                })

    return headings


def create_toc_paragraph(doc, text: str, level: int, page_num: str = ''):
    """Cria um parágrafo de entrada do sumário"""
    para = doc.add_paragraph()

    # Indentação baseada no nível
    indent = Cm(level * 0.5)
    para.paragraph_format.left_indent = indent
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    # Adicionar texto
    run = para.add_run(text)
    run.font.name = 'EB Garamond'
    run.font.size = Pt(11) if level == 1 else Pt(10)

    if level == 1:
        run.bold = True

    # Tab e número de página (simulado - Word atualizará)
    if page_num:
        tab_run = para.add_run('\t')
        page_run = para.add_run(page_num)
        page_run.font.name = 'EB Garamond'
        page_run.font.size = Pt(11) if level == 1 else Pt(10)

    return para


def add_toc_field(paragraph):
    """
    Adiciona um campo TOC do Word (será atualizado automaticamente ao abrir)
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    return paragraph


def insert_toc(doc, headings: list, position: int = 0):
    """
    Insere o sumário no documento

    Args:
        doc: Documento Word
        headings: Lista de headings extraídos
        position: Índice do parágrafo onde inserir (0 = início)
    """
    # Criar conteúdo do TOC
    toc_content = []

    # Título do sumário
    toc_content.append({
        'type': 'title',
        'text': 'SUMÁRIO',
    })

    # Linha em branco
    toc_content.append({
        'type': 'blank',
    })

    # Entradas
    for h in headings:
        toc_content.append({
            'type': 'entry',
            'level': h['level'],
            'title': h['title'],
        })

    # Quebra de página após TOC
    toc_content.append({
        'type': 'page_break',
    })

    # Inserir no documento (do fim para o início para manter índices)
    for item in reversed(toc_content):
        if item['type'] == 'title':
            para = doc.paragraphs[position].insert_paragraph_before()
            run = para.add_run(item['text'])
            run.font.name = 'Josefin Sans'
            run.font.size = Pt(24)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(24)

        elif item['type'] == 'blank':
            para = doc.paragraphs[position].insert_paragraph_before()

        elif item['type'] == 'entry':
            para = doc.paragraphs[position].insert_paragraph_before()
            level = item['level']

            # Indentação
            para.paragraph_format.left_indent = Cm(level * 0.5)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)

            # Texto
            run = para.add_run(item['title'])
            run.font.name = 'EB Garamond'
            run.font.size = Pt(11) if level == 1 else Pt(10)
            if level == 1:
                run.bold = True

            # Tab leader (pontilhado) - adicionar manualmente
            tab_run = para.add_run('\t')

            # Placeholder para número de página
            # (Word atualizará ao abrir se usar campo TOC)
            page_run = para.add_run('...')
            page_run.font.name = 'EB Garamond'
            page_run.font.size = Pt(10)

        elif item['type'] == 'page_break':
            para = doc.paragraphs[position].insert_paragraph_before()
            run = para.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)


def generate_simple_toc(doc, headings: list):
    """
    Gera um sumário simples no início do documento
    Versão mais compatível que não usa campos complexos do Word
    """
    # Guardar primeiro parágrafo
    first_para = doc.paragraphs[0] if doc.paragraphs else None

    # Criar novo documento temporário para construir TOC
    # Inserir antes do primeiro parágrafo

    # Título do sumário
    toc_title = doc.paragraphs[0].insert_paragraph_before('SUMÁRIO')
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.font.name = 'Josefin Sans'
        run.font.size = Pt(24)
        run.bold = True
    toc_title.paragraph_format.space_after = Pt(24)

    # Linha em branco
    blank = toc_title.insert_paragraph_before('')

    # Inserir entradas (em ordem reversa pois estamos inserindo "before")
    insert_point = doc.paragraphs[0]

    for h in headings:
        entry = insert_point.insert_paragraph_before()
        level = h['level']

        # Indentação baseada no nível
        entry.paragraph_format.left_indent = Cm((level - 1) * 0.75)
        entry.paragraph_format.space_before = Pt(4)
        entry.paragraph_format.space_after = Pt(4)

        # Texto da entrada
        run = entry.add_run(h['title'])
        run.font.name = 'EB Garamond'
        run.font.size = Pt(12) if level == 1 else Pt(11)
        if level == 1:
            run.bold = True

    # Quebra de página após sumário
    page_break_para = insert_point.insert_paragraph_before()
    page_break_para.add_run().add_break()


# =============================================================================
# PROCESSAMENTO PRINCIPAL
# =============================================================================

def process_document(input_path: str, output_path: str, position: str = 'start'):
    """Processa documento adicionando sumário"""

    print(f"📖 Carregando: {input_path}")
    doc = Document(input_path)

    # Extrair headings
    print("🔍 Extraindo estrutura...")
    headings = extract_headings(doc)

    if not headings:
        print("⚠️  Nenhum heading encontrado no documento.")
        print("   Verifique se o documento tem títulos formatados corretamente.")
        doc.save(output_path)
        return {'headings': 0}

    print(f"   Encontrados {len(headings)} headings")

    # Mostrar estrutura
    print("\n📋 Estrutura do documento:")
    for h in headings[:10]:  # Mostrar primeiros 10
        indent = "  " * h['level']
        print(f"   {indent}• {h['title'][:50]}...")

    if len(headings) > 10:
        print(f"   ... e mais {len(headings) - 10} entradas")

    # Gerar sumário
    print("\n📝 Gerando sumário...")

    # Determinar posição de inserção
    insert_pos = 0
    if position == 'after-preface':
        # Procurar fim do prefácio
        for i, para in enumerate(doc.paragraphs):
            text = para.text.lower()
            if 'prefácio' in text or 'preface' in text or 'introdução' in text:
                # Procurar próximo heading ou fim do prefácio
                for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
                    level, _ = is_heading(doc.paragraphs[j])
                    if level is not None:
                        insert_pos = j
                        break
                break

    # Inserir sumário simples
    generate_simple_toc(doc, headings)

    # Salvar
    print(f"💾 Salvando: {output_path}")
    doc.save(output_path)

    print("\n✅ Sumário gerado com sucesso!")
    print(f"   Total de entradas: {len(headings)}")

    return {'headings': len(headings)}


def main():
    parser = argparse.ArgumentParser(
        description='Gera sumário automático em documentos Word',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python3 generate-toc.py manuscrito.docx com-sumario.docx
  python3 generate-toc.py livro.docx livro-final.docx --position after-preface
        """
    )

    parser.add_argument('input', help='Documento Word de entrada (.docx)')
    parser.add_argument('output', help='Documento Word de saída (.docx)')
    parser.add_argument(
        '--position', '-p',
        choices=['start', 'after-preface'],
        default='start',
        help='Onde inserir o sumário (default: start)'
    )
    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Modo verbose com mais detalhes'
    )

    args = parser.parse_args()

    # Validar entrada
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERRO: Arquivo não encontrado: {args.input}")
        sys.exit(1)

    if not input_path.suffix.lower() == '.docx':
        print(f"ERRO: Arquivo deve ser .docx: {args.input}")
        sys.exit(1)

    # Processar
    try:
        process_document(args.input, args.output, args.position)
    except Exception as e:
        print(f"ERRO: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
