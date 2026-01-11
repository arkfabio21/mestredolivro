#!/usr/bin/env python3
"""
quality-guard.py - Detecta e elimina vícios de escrita de IA em textos

USAGE:
    python3 scripts/quality-guard.py input.md output.md [--fix] [--report]
    python3 scripts/quality-guard.py input.docx output.docx [--fix] [--report]

DESCRIÇÃO:
    Analisa texto e detecta padrões que denunciam escrita gerada por IA:
    - Pontuação artificial (travessões, dois pontos excessivos)
    - Transições clichê (Além disso, Portanto, No entanto)
    - Estruturas artificiais (listas de 3, frases curtas)
    - Frases genéricas (É fundamental, Cabe ressaltar)

MODOS:
    --analyze: Apenas analisar e reportar (default)
    --fix: Aplicar correções automáticas
    --report: Gerar relatório detalhado

DEPENDÊNCIAS:
    pip install python-docx
"""

import argparse
import sys
import re
import json
from pathlib import Path
from collections import Counter
from typing import List, Dict, Tuple

try:
    from docx import Document
except ImportError:
    Document = None  # Suporte a DOCX opcional


# =============================================================================
# PADRÕES DE VÍCIOS DE IA
# =============================================================================

# Transições clichê
TRANSICOES_CLICHE = [
    r'\bAlém disso\b',
    r'\bPortanto\b',
    r'\bNo entanto\b',
    r'\bEm resumo\b',
    r'\bÉ importante ressaltar\b',
    r'\bVale destacar que\b',
    r'\bNesse sentido\b',
    r'\bDiante do exposto\b',
    r'\bCabe ressaltar que\b',
    r'\bPor fim\b',
    r'\bEm conclusão\b',
    r'\bDessa forma\b',
    r'\bAssim sendo\b',
    r'\bÉ fundamental\b',
    r'\bTorna-se relevante\b',
    r'\bPodemos concluir que\b',
    r'\bÉ válido mencionar\b',
    r'\bConvém lembrar que\b',
    r'\bNota-se que\b',
    r'\bPercebe-se que\b',
]

# Frases genéricas
FRASES_GENERICAS = [
    r'É fundamental compreender a importância de',
    r'Nesse contexto, torna-se relevante observar que',
    r'Diante do exposto, podemos concluir que',
    r'é de suma importância',
    r'desempenha um papel crucial',
    r'desempenha um papel fundamental',
    r'é essencial para o sucesso',
    r'podemos observar que',
    r'é possível afirmar que',
    r'ao longo dos anos',
    r'nos dias de hoje',
    r'desde os primórdios',
]

# Estruturas artificiais (numeração)
ESTRUTURAS_NUMERADAS = [
    r'\bPrimeiro\b.*\bSegundo\b.*\bTerceiro\b',
    r'\bEm primeiro lugar\b',
    r'\bEm segundo lugar\b',
    r'\bEm terceiro lugar\b',
    r'\bPrimeiramente\b',
]

# Padrões de substituição
SUBSTITUICOES = {
    r'—': ', ',  # Travessão → vírgula
    r'Além disso, ': '',  # Remover transição
    r'Portanto, ': '',
    r'No entanto, ': 'Mas ',
    r'Em resumo, ': '',
    r'É importante ressaltar que ': '',
    r'Vale destacar que ': '',
    r'Nesse sentido, ': '',
    r'Diante do exposto, ': '',
    r'Cabe ressaltar que ': '',
    r'É fundamental ': '',
    r'Dessa forma, ': '',
    r'Assim sendo, ': '',
}


# =============================================================================
# CLASSE PRINCIPAL
# =============================================================================

class QualityGuard:
    """Analisador de vícios de escrita de IA"""

    def __init__(self, text: str):
        self.original_text = text
        self.text = text
        self.issues = []
        self.stats = Counter()

    def analyze(self) -> Dict:
        """Analisa o texto e retorna relatório de problemas"""
        self.issues = []
        self.stats = Counter()

        # Dividir em parágrafos
        paragraphs = [p.strip() for p in self.text.split('\n\n') if p.strip()]

        # Analisar densidade de headings (vício de IA: muitos subtítulos)
        self._analyze_heading_density(paragraphs)

        # Detectar tabelas Markdown
        self._analyze_markdown_tables(paragraphs)

        for i, para in enumerate(paragraphs, 1):
            self._analyze_paragraph(para, i)

        return self._generate_report()

    def _analyze_heading_density(self, paragraphs: list):
        """Analisa subtítulos vazios (com pouco conteúdo abaixo)"""
        # Encontrar posições dos subtítulos (## apenas, ignorar título principal #)
        heading_positions = []
        for i, para in enumerate(paragraphs):
            stripped = para.strip()
            # Apenas ## (subtítulos), não # (título principal do capítulo)
            if stripped.startswith('##') and not stripped.startswith('###'):
                heading_positions.append(i)
            elif stripped.startswith('###'):
                heading_positions.append(i)

        # Verificar conteúdo entre subtítulos
        subtitulos_vazios = 0
        for j, pos in enumerate(heading_positions):
            # Calcular próxima posição (próximo heading ou fim)
            next_pos = heading_positions[j + 1] if j + 1 < len(heading_positions) else len(paragraphs)

            # Contar parágrafos de conteúdo entre este heading e o próximo
            content_count = 0
            content_words = 0
            for k in range(pos + 1, next_pos):
                para = paragraphs[k]
                if not para.strip().startswith('#') and len(para.strip()) > 0:
                    content_count += 1
                    content_words += len(para.split())

            # Subtítulo "vazio" = menos de 2 parágrafos OU menos de 100 palavras
            if content_count < 2 or content_words < 100:
                subtitulos_vazios += 1
                self.issues.append({
                    'tipo': 'subtitulo_vazio',
                    'paragrafo': pos,
                    'severidade': 'media',
                    'descricao': f'Subtítulo com pouco conteúdo ({content_count} parágrafos, {content_words} palavras)',
                })

        if subtitulos_vazios > 0:
            self.stats['subtitulo_vazio'] = subtitulos_vazios

    def _analyze_markdown_tables(self, paragraphs: list):
        """Detecta tabelas Markdown no texto"""
        for i, para in enumerate(paragraphs, 1):
            # Detectar linhas de tabela Markdown
            if re.search(r'\|.*\|.*\|', para):
                self.issues.append({
                    'tipo': 'tabela_markdown',
                    'paragrafo': i,
                    'severidade': 'critica',
                    'descricao': 'Tabela Markdown detectada - converter para Excel',
                })
                self.stats['tabela_markdown'] += 1

    def _analyze_paragraph(self, para: str, para_num: int):
        """Analisa um parágrafo individual"""

        # 1. Travessões
        count = para.count('—')
        if count > 0:
            self.issues.append({
                'tipo': 'travessao',
                'paragrafo': para_num,
                'quantidade': count,
                'severidade': 'alta',
                'descricao': f'{count} travessão(ões) encontrado(s)',
            })
            self.stats['travessao'] += count

        # 2. Dois pontos excessivos
        colons = para.count(':')
        if colons > 2:
            self.issues.append({
                'tipo': 'dois_pontos',
                'paragrafo': para_num,
                'quantidade': colons,
                'severidade': 'media',
                'descricao': f'{colons} dois pontos (máximo recomendado: 2)',
            })
            self.stats['dois_pontos'] += colons

        # 3. Ponto e vírgula
        semicolons = para.count(';')
        if semicolons > 1:
            self.issues.append({
                'tipo': 'ponto_virgula',
                'paragrafo': para_num,
                'quantidade': semicolons,
                'severidade': 'baixa',
                'descricao': f'{semicolons} ponto e vírgula (máximo recomendado: 1)',
            })
            self.stats['ponto_virgula'] += semicolons

        # 4. Transições clichê
        for pattern in TRANSICOES_CLICHE:
            matches = re.findall(pattern, para, re.IGNORECASE)
            if matches:
                for match in matches:
                    self.issues.append({
                        'tipo': 'transicao_cliche',
                        'paragrafo': para_num,
                        'texto': match,
                        'severidade': 'alta',
                        'descricao': f'Transição clichê: "{match}"',
                    })
                    self.stats['transicao_cliche'] += 1

        # 5. Frases genéricas
        for pattern in FRASES_GENERICAS:
            if re.search(pattern, para, re.IGNORECASE):
                self.issues.append({
                    'tipo': 'frase_generica',
                    'paragrafo': para_num,
                    'severidade': 'alta',
                    'descricao': f'Frase genérica detectada',
                })
                self.stats['frase_generica'] += 1

        # 6. Frases muito curtas (menos de 8 palavras isoladas)
        sentences = re.split(r'[.!?]', para)
        short_count = 0
        for sent in sentences:
            words = sent.split()
            if 2 <= len(words) < 8:
                short_count += 1

        if short_count >= 3:
            self.issues.append({
                'tipo': 'frases_curtas',
                'paragrafo': para_num,
                'quantidade': short_count,
                'severidade': 'media',
                'descricao': f'{short_count} frases curtas consecutivas',
            })
            self.stats['frases_curtas'] += short_count

        # 7. Estruturas numeradas artificiais
        for pattern in ESTRUTURAS_NUMERADAS:
            if re.search(pattern, para, re.IGNORECASE):
                self.issues.append({
                    'tipo': 'estrutura_numerada',
                    'paragrafo': para_num,
                    'severidade': 'media',
                    'descricao': 'Estrutura numerada artificial (Primeiro, Segundo...)',
                })
                self.stats['estrutura_numerada'] += 1
                break

    def fix(self) -> str:
        """Aplica correções automáticas ao texto"""
        fixed = self.text

        # Aplicar substituições
        for pattern, replacement in SUBSTITUICOES.items():
            fixed = re.sub(pattern, replacement, fixed, flags=re.IGNORECASE)

        # Limpar espaços duplos resultantes
        fixed = re.sub(r'  +', ' ', fixed)
        fixed = re.sub(r'\n +', '\n', fixed)

        # Capitalizar início de frases após remoções
        fixed = re.sub(r'\.\s+([a-z])', lambda m: '. ' + m.group(1).upper(), fixed)

        self.text = fixed
        return fixed

    def _generate_report(self) -> Dict:
        """Gera relatório estruturado"""
        total = sum(self.stats.values())

        # Agrupar por tipo
        by_type = {}
        for issue in self.issues:
            tipo = issue['tipo']
            if tipo not in by_type:
                by_type[tipo] = []
            by_type[tipo].append(issue)

        return {
            'total_vicios': total,
            'por_tipo': dict(self.stats),
            'detalhes': by_type,
            'issues': self.issues,
            'status': 'CLEANED' if total == 0 else 'NEEDS_REVIEW',
        }

    def get_markdown_report(self) -> str:
        """Gera relatório em formato Markdown"""
        report = self._generate_report()

        lines = [
            '# RELATÓRIO QUALITY GUARD',
            '',
            f'**Total de vícios encontrados:** {report["total_vicios"]}',
            f'**Status:** {report["status"]}',
            '',
            '## Resumo por Tipo',
            '',
            '| Tipo | Quantidade |',
            '|------|------------|',
        ]

        type_names = {
            'travessao': 'Travessões',
            'dois_pontos': 'Dois pontos excessivos',
            'ponto_virgula': 'Ponto e vírgula',
            'transicao_cliche': 'Transições clichê',
            'frase_generica': 'Frases genéricas',
            'frases_curtas': 'Frases curtas',
            'estrutura_numerada': 'Estruturas numeradas',
            'subtitulo_vazio': 'Subtítulos vazios',
            'tabela_markdown': 'Tabelas Markdown',
        }

        for tipo, qtd in report['por_tipo'].items():
            nome = type_names.get(tipo, tipo)
            lines.append(f'| {nome} | {qtd} |')

        if report['issues']:
            lines.extend([
                '',
                '## Detalhes',
                '',
            ])

            current_para = None
            for issue in sorted(report['issues'], key=lambda x: x['paragrafo']):
                if issue['paragrafo'] != current_para:
                    current_para = issue['paragrafo']
                    lines.append(f'### Parágrafo {current_para}')
                    lines.append('')

                lines.append(f'- **{type_names.get(issue["tipo"], issue["tipo"])}**: {issue["descricao"]}')

            lines.append('')

        lines.extend([
            '## Checklist de Varredura',
            '',
            f'- [{"x" if report["por_tipo"].get("travessao", 0) == 0 else " "}] Zero travessões (—)',
            f'- [{"x" if report["por_tipo"].get("dois_pontos", 0) <= 2 else " "}] Dois pontos: máximo 2 por página',
            f'- [{"x" if report["por_tipo"].get("ponto_virgula", 0) <= 1 else " "}] Ponto e vírgula: máximo 1 por página',
            f'- [{"x" if report["por_tipo"].get("frases_curtas", 0) == 0 else " "}] Zero frases curtas isoladas',
            f'- [{"x" if report["por_tipo"].get("transicao_cliche", 0) == 0 else " "}] Zero transições clichê',
            f'- [{"x" if report["por_tipo"].get("frase_generica", 0) == 0 else " "}] Zero frases genéricas',
            f'- [{"x" if report["por_tipo"].get("estrutura_numerada", 0) == 0 else " "}] Zero estruturas numeradas artificiais',
            f'- [{"x" if report["por_tipo"].get("subtitulo_vazio", 0) == 0 else " "}] Subtítulos com conteúdo substancial',
            f'- [{"x" if report["por_tipo"].get("tabela_markdown", 0) == 0 else " "}] Zero tabelas Markdown (usar Excel)',
        ])

        return '\n'.join(lines)


# =============================================================================
# FUNÇÕES DE I/O
# =============================================================================

def read_file(path: str) -> str:
    """Lê arquivo (Markdown ou DOCX)"""
    path = Path(path)

    if path.suffix.lower() == '.docx':
        if Document is None:
            print("ERRO: python-docx não instalado para ler .docx")
            print("Execute: pip install python-docx")
            sys.exit(1)

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        return '\n\n'.join(paragraphs)

    else:
        return path.read_text(encoding='utf-8')


def write_file(path: str, content: str, original_path: str = None):
    """Escreve arquivo (Markdown ou DOCX)"""
    path = Path(path)

    if path.suffix.lower() == '.docx':
        if Document is None:
            print("ERRO: python-docx não instalado para escrever .docx")
            sys.exit(1)

        # Se temos original, usar como base
        if original_path and Path(original_path).suffix.lower() == '.docx':
            doc = Document(original_path)
            # Substituir texto dos parágrafos
            paragraphs = content.split('\n\n')
            for i, para in enumerate(doc.paragraphs):
                if i < len(paragraphs):
                    para.text = paragraphs[i]
        else:
            doc = Document()
            for para in content.split('\n\n'):
                doc.add_paragraph(para)

        doc.save(path)
    else:
        path.write_text(content, encoding='utf-8')


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Detecta e elimina vícios de escrita de IA',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python3 quality-guard.py capitulo.md --report
  python3 quality-guard.py capitulo.md capitulo-limpo.md --fix
  python3 quality-guard.py documento.docx documento-limpo.docx --fix --report
        """
    )

    parser.add_argument('input', help='Arquivo de entrada (.md ou .docx)')
    parser.add_argument('output', nargs='?', help='Arquivo de saída (opcional)')
    parser.add_argument(
        '--fix', '-f',
        action='store_true',
        help='Aplicar correções automáticas'
    )
    parser.add_argument(
        '--report', '-r',
        action='store_true',
        help='Gerar relatório detalhado'
    )
    parser.add_argument(
        '--json',
        action='store_true',
        help='Saída em formato JSON'
    )
    parser.add_argument(
        '--quiet', '-q',
        action='store_true',
        help='Saída mínima'
    )

    args = parser.parse_args()

    # Validar entrada
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERRO: Arquivo não encontrado: {args.input}")
        sys.exit(1)

    # Ler arquivo
    if not args.quiet:
        print(f"📖 Analisando: {args.input}")

    text = read_file(args.input)
    guard = QualityGuard(text)

    # Analisar
    report = guard.analyze()

    # Exibir resultado
    if args.json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    elif args.report:
        print(guard.get_markdown_report())
    elif not args.quiet:
        print(f"\n📊 Resultado:")
        print(f"   Vícios encontrados: {report['total_vicios']}")
        print(f"   Status: {report['status']}")

        if report['por_tipo']:
            print(f"\n   Por tipo:")
            for tipo, qtd in report['por_tipo'].items():
                print(f"   - {tipo}: {qtd}")

    # Aplicar correções se solicitado
    if args.fix:
        if not args.quiet:
            print(f"\n✨ Aplicando correções...")

        fixed_text = guard.fix()

        # Re-analisar
        guard2 = QualityGuard(fixed_text)
        report2 = guard2.analyze()

        if not args.quiet:
            print(f"   Vícios restantes: {report2['total_vicios']}")

        # Salvar
        if args.output:
            write_file(args.output, fixed_text, args.input)
            if not args.quiet:
                print(f"\n💾 Salvo em: {args.output}")
        else:
            # Sem arquivo de saída, mostrar texto corrigido
            print("\n--- TEXTO CORRIGIDO ---")
            print(fixed_text)
            print("--- FIM ---")

    # Código de saída
    if report['total_vicios'] > 0 and not args.fix:
        sys.exit(1)  # Indicar que há problemas
    sys.exit(0)


if __name__ == '__main__':
    main()
